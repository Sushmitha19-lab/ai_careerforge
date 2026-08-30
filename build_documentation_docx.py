from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


TEAL = RGBColor(0x35, 0x6B, 0x72)
DARK = RGBColor(0x17, 0x21, 0x2B)
MUTED = RGBColor(0x5D, 0x65, 0x6A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CORAL = RGBColor(0xE8, 0x75, 0x5F)


def shade(cell, hex_color):
    fill = hex_color.lstrip("#")
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def borders(table, color="C9C4BB"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders_el = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders_el.append(el)
    tbl_pr.append(borders_el)


def set_run(run, size=11, bold=False, color=DARK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    r = run._element
    r_pr = r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font)
    r_fonts.set(qn("w:hAnsi"), font)
    r_fonts.set(qn("w:eastAsia"), font)


def add_heading_custom(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    if level == 1:
        set_run(run, size=18, bold=True, color=TEAL)
    elif level == 2:
        set_run(run, size=14, bold=True, color=DARK)
    else:
        set_run(run, size=12, bold=True, color=TEAL)
    return p


def add_body(doc, text, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=11, color=DARK)
    return p


def add_mixed(doc, parts, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    for text, bold in parts:
        run = p.add_run(text)
        set_run(run, size=11, bold=bold, color=DARK)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run(r1, size=11, bold=True, color=DARK)
        r2 = p.add_run(text)
        set_run(r2, size=11, color=DARK)
    else:
        r = p.add_run(text)
        set_run(r, size=11, color=DARK)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=11, color=DARK)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_run(run, size=9, bold=True, color=MUTED)
    run.italic = True
    return p


def add_figure_image(doc, image_path, width_in=6.5):
    doc.add_picture(str(image_path), width=Inches(width_in))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_visual(doc, png_path, ascii_lines, caption):
    if png_path and Path(png_path).exists():
        add_figure_image(doc, png_path)
    else:
        add_diagram(doc, ascii_lines)
    add_caption(doc, caption)


def add_diagram(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.2)
    run = p.add_run("\n".join(lines))
    set_run(run, size=8, color=DARK, font="Consolas")
    shade_p = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F3EE")
    shade_p.append(shd)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    set_run(run, size=9.5, color=DARK, font="Consolas")
    shade_p = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F3EE")
    shade_p.append(shd)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    borders(table)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        set_run(run, size=10, bold=True, color=WHITE)
        shade(cell, "356B72")

    for r_idx, row in enumerate(rows):
        fill = "FFFFFF" if r_idx % 2 == 0 else "F5F3EE"
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            set_run(run, size=10, color=DARK)
            shade(cell, fill)
            if col_widths:
                cell.width = Cm(col_widths[c_idx])

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def add_sequence_table(doc, steps):
    add_table(
        doc,
        ["Step", "From", "To", "Action"],
        [[str(i + 1), a, b, c] for i, (a, b, c) in enumerate(steps)],
        col_widths=[1.6, 3.4, 3.6, 8.2],
    )


def build():
    figure_paths = {}
    try:
        from build_architecture_figures import main as draw_figures

        figure_paths = draw_figures()
    except Exception as exc:
        print("Skipping rendered figures:", exc)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("CareerForge  |  AI Interview Career Support")
    set_run(hr, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Project documentation  ·  Confidential for academic use")
    set_run(fr, size=8, color=MUTED)

    # Cover
    for _ in range(4):
        doc.add_paragraph()

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("PROJECT DOCUMENTATION")
    set_run(kr, size=12, bold=True, color=CORAL)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    tr = title.add_run("CareerForge")
    set_run(tr, size=36, bold=True, color=TEAL)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("AI Interview Career Support System")
    set_run(sr, size=16, color=DARK)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = line.add_run("Architecture  ·  Block diagrams  ·  Sequence  ·  Usage")
    set_run(lr, size=11, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(28)
    mr = meta.add_run(
        "Web application for company-specific mock interviews\n"
        "with keyword scoring, camera behavior analysis, voice fluency and accent ML,\n"
        "SkillBridge maps, and a 10-course / 10-company practice catalog"
    )
    set_run(mr, size=11, color=DARK)

    stack = doc.add_paragraph()
    stack.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stack.paragraph_format.space_before = Pt(36)
    st = stack.add_run("React  ·  Express  ·  MySQL  ·  Flask  ·  OpenCV  ·  scikit-learn")
    set_run(st, size=10, bold=True, color=TEAL)

    doc.add_page_break()

    # TOC-like contents
    add_heading_custom(doc, "Contents", 1)
    contents = [
        "1. Introduction",
        "2. System architecture",
        "    2.1 Architecture diagram",
        "    2.2 Product block diagram",
        "    2.3 Scoring block diagram",
        "    2.4 Components and page flow",
        "    2.5 ML modules",
        "3. Sequence flows",
        "4. Usage (run on another computer)",
        "    4.1 Software to install",
        "    4.2 What to copy",
        "    4.3 Windows commands",
        "    4.4 macOS / Linux commands",
        "    4.5 Optional MySQL",
        "    4.6 Optional Google Sign-In",
        "    4.7 Check it is running",
        "    4.8 How to use the product",
        "    4.9 API map",
        "5. Pros and cons",
        "6. Project structure",
        "7. Limitations and future work",
        "8. Summary",
    ]
    for item in contents:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run(r, size=12, color=DARK)

    # 1
    add_heading_custom(doc, "1. Introduction", 1)
    add_body(
        doc,
        "CareerForge is a web application that helps students prepare for company-specific "
        "technical interviews. The student signs in (email/password or Google), works through a "
        "10-course × 10-company catalog and a Practice / SkillBridge workspace, then takes a "
        "four-round mock interview. During each round the system evaluates what they say and "
        "how they present themselves.",
    )

    add_table(
        doc,
        ["Signal", "What it measures"],
        [
            ["Answer quality", "Keyword matching and pattern matching on the typed or spoken answer"],
            ["Camera behavior", "Confidence, nervousness, and malpractice risk from the webcam"],
            ["Voice delivery", "Fluency, content accuracy, and accent/clarity (ASR intelligibility)"],
            ["Skill / chance", "Technical, communication, problem-solving, plus chance of advancing this screen"],
        ],
        col_widths=[4.2, 12.6],
    )
    add_caption(doc, "Table 1. Evaluation signals used in the mock interview")

    add_body(
        doc,
        "The goal is structured practice: company-specific rounds, adaptive difficulty across four "
        "rounds, camera and voice signals, and a results pack after submit. Rubric chips are not "
        "shown live in the answer box.",
    )

    add_heading_custom(doc, "1.1 Problem", 2)
    add_body(
        doc,
        "Students often know the syllabus but freeze in interviews. Typical prep apps score only "
        "written answers, or only length. They miss:",
    )
    add_bullet(doc, "whether the answer actually contains the expected concepts")
    add_bullet(doc, "whether the candidate looks away, leaves the frame, or appears with another person")
    add_bullet(doc, "whether spoken answers are fluent and on-topic")

    add_heading_custom(doc, "1.2 Objectives", 2)
    add_numbered(doc, "Guide the student from login to dashboard, practice or course/company, four-round mock, and result.")
    add_numbered(doc, "Score answers with keywords (concepts) and patterns (definition, comparison, example).")
    add_numbered(doc, "Analyze the webcam for confidence, nervousness, and integrity flags.")
    add_numbered(doc, "Score voice answers for fluency, accuracy, and accent/clarity (ASR intelligibility, not a linguistic accent classifier).")
    add_numbered(doc, "Keep the ML stack in a separate Python service so models can be retrained without rewriting the UI.")
    add_numbered(doc, "Persist accounts and sessions in MySQL when reachable, otherwise JSON files under backend/data/.")

    # 2
    add_heading_custom(doc, "2. System architecture", 1)
    add_mixed(
        doc,
        [
            ("CareerForge is a three-process system. ", False),
            (
                "The browser never talks to MySQL, JSON account files, or .pkl / ONNX models directly. "
                "Vite (and Express after a production build) proxy /api and /ml so the student uses one origin.",
                False,
            ),
        ],
    )

    add_heading_custom(doc, "2.1 Architecture diagram", 2)
    add_body(
        doc,
        "Figure 1 is the deployment architecture: student browser, Express API, Flask ML service, "
        "optional Google Identity, and persistence that prefers MySQL and falls back to JSON files.",
    )
    add_visual(
        doc,
        figure_paths.get("architecture"),
        [
            "                         STUDENT BROWSER  (React + Vite :5173)",
            "  +------------------+   +------------------+   +------------------+",
            "  | Pages / Router   |   | Keyword matcher  |   | Camera + Mic     |",
            "  | login, practice, |   | evaluateAnswer   |   | JPEG + transcript|",
            "  | mock, SkillBridge|   |                  |   |                  |",
            "  +--------+---------+   +------------------+   +--------+---------+",
            "           |  /api/*                                          |  /ml/*",
            "           v                                                  v",
            "  +------------------------------+            +------------------------------+",
            "  | Express API :5000            |            | Flask ML :5001               |",
            "  | JWT + Google ID-token verify |            | POST /analyze-emotion        |",
            "  | Catalog 10 courses x 10 cos. |            | POST /analyze-voice          |",
            "  | Save interview sessions      |            | YuNet + voice regressor      |",
            "  +--------------+---------------+            +------------------------------+",
            "                 |",
            "         +-------+--------+",
            "         |                |",
            "         v                v",
            "  +-------------+  +-------------------+",
            "  | MySQL       |  | JSON file store   |",
            "  | if reachable|  | if MySQL times out|",
            "  +-------------+  +-------------------+",
        ],
        "Figure 1. System architecture diagram",
    )

    add_table(
        doc,
        ["Block", "Runs on", "Talks to", "Does not talk to"],
        [
            ["React SPA", "Browser :5173", "Express /api, Flask /ml, Google GIS", "MySQL, .pkl files, ONNX"],
            ["Express", "Node :5000", "MySQL or JSON files, Google tokeninfo", "YuNet / sklearn directly"],
            ["Flask ML", "Python :5001", "ONNX + pickle models on disk", "The student browser origin except via /ml proxy"],
            ["Google Identity", "accounts.google.com", "Returns an ID token to the login page", "CareerForge database"],
        ],
        col_widths=[3.2, 3.4, 5.2, 5.0],
    )
    add_caption(doc, "Table 2. Who may call whom")

    add_heading_custom(doc, "2.2 Product block diagram", 2)
    add_body(
        doc,
        "Figure 2 is the product block diagram (what the student uses), not the process diagram. "
        "Prepare feeds a four-round mock. Three scoring blocks feed the result pack.",
    )
    add_visual(
        doc,
        figure_paths.get("product"),
        [
            "+------------------+      +---------------------------+      +--------------------+",
            "| PREPARE          |      | MOCK (4 adaptive rounds)  |      | OUTPUT             |",
            "|                  |      |                           |      |                    |",
            "| 10 courses       |      | 1  Aptitude               |      | Company readiness  |",
            "| 10 companies     |----->| 2  Coding                 |----->| Chance of advancing|",
            "| Round resources  |      | 3  Technical              |      | Weak areas         |",
            "| SkillBridge map  |      | 4  HR                     |      | Saved to account   |",
            "| Week-wise coding |      |                           |      |                    |",
            "+------------------+      +-------------+-------------+      +--------------------+",
            "                                        |",
            "                          +-------------+-------------+",
            "                          | SCORING BLOCKS            |",
            "                          |  Text  |  Camera  | Voice |",
            "                          +---------------------------+",
        ],
        "Figure 2. Product block diagram",
    )

    add_table(
        doc,
        ["Block", "What it contains", "Student sees it on"],
        [
            ["Prepare", "10x10 catalog, course-specific rounds, live practice URLs, SkillBridge, week-wise coding", "/practice, /readiness, /preparation"],
            ["Mock", "Four adaptive rounds: aptitude, coding, technical, HR. Difficulty follows the last score.", "/interview"],
            ["Scoring", "Keywords+patterns; YuNet camera; voice fluency, accuracy, accent/clarity", "Interview meters + /result"],
            ["Output", "Readiness %, chance of advancing this screen, weak areas, saved session", "/result, account history"],
        ],
        col_widths=[3.0, 8.4, 5.4],
    )
    add_caption(doc, "Table 3. Product blocks")

    add_heading_custom(doc, "2.3 Scoring block diagram", 2)
    add_body(
        doc,
        "Each submitted round runs three independent blocks. The live answer box does not show the "
        "rubric; scores appear after submit.",
    )
    add_visual(
        doc,
        figure_paths.get("scoring"),
        [
            "                    +---------------- Student answer ----------------+",
            "                    |  typed text and/or spoken transcript           |",
            "                    +-----------+------------------+-----------------+",
            "                                |                  |",
            "               BLOCK A          |                  |         BLOCK C",
            "          Answer quality        |                  |           Voice",
            "   +----------------------+     |     BLOCK B      |    +----------------------+",
            "   | Keyword groups       |     |     Camera       |    | Web Speech transcript|",
            "   | Regex patterns       |     | +--------------+ |    | WPM, fillers, conf.  |",
            "   | Length completeness  |     | | JPEG ~1.4s   | |    | RF/GBR or heuristic  |",
            "   |                      |     | | YuNet+points | |    |                      |",
            "   | 50% + 35% + 15%      |     | | integrity    | |    | fluency, accuracy,   |",
            "   | = round text score   |     | +--------------+ |    | accent/clarity       |",
            "   +----------+-----------+     +--------+---------+    +----------+-----------+",
            "              |                          |                         |",
            "              +-------------+------------+-------------+-----------+",
            "                            v",
            "              +------------------------------------------+",
            "              | Result pack + explainInterview()         |",
            "              | readiness, chance, SkillBridge scores    |",
            "              +------------------------------------------+",
        ],
        "Figure 3. Interview scoring block diagram",
    )

    add_heading_custom(doc, "2.4 Components and page flow", 2)
    add_table(
        doc,
        ["Layer", "Technology", "Port", "Responsibility"],
        [
            ["Frontend", "React + Vite + React Router", "5173", "Login, practice, catalog, mock, result, SkillBridge"],
            ["Backend API", "Express.js", "5000", "JWT auth, Google verify, catalog, saved interviews"],
            ["Database", "MySQL (optional)", "3306", "Users and interview sessions when reachable"],
            ["File store", "JSON in backend/data/", "—", "Auth and results if MySQL times out"],
            ["ML service", "Flask + OpenCV + sklearn", "5001", "Camera integrity, voice fluency / accuracy / accent"],
        ],
        col_widths=[3.0, 4.6, 2.2, 7.0],
    )
    add_caption(doc, "Table 4. System layers")

    add_body(doc, "Routing is React Router in src/App.jsx.")
    add_table(
        doc,
        ["From", "To", "Trigger"],
        [
            ["Login", "Dashboard", "Email/password or Google sign-in"],
            ["Dashboard", "Practice", "Practice workspace"],
            ["Dashboard", "Courses", "Mock path — choose a domain"],
            ["Dashboard", "SkillBridge /readiness", "Maps, week-wise coding, requirement detector"],
            ["Courses", "Companies", "Course chosen (10 companies)"],
            ["Companies / Practice", "Preparation or interview", "Company chosen"],
            ["Mock interview", "Result", "Four rounds complete or End interview"],
            ["Result", "Dashboard, retry, or SkillBridge", "After scores and chance %"],
        ],
        col_widths=[4.5, 5.5, 7.0],
    )
    add_caption(doc, "Figure 4. User navigation flow")

    add_heading_custom(doc, "2.5 ML modules", 2)
    add_mixed(
        doc,
        [
            ("Skill model (train_model.py): ", True),
            (
                "RandomForestClassifier on technical_score, communication_score, "
                "problem_solving maps to skill_level (Strong / Average / Weak).",
                False,
            ),
        ],
    )
    add_mixed(
        doc,
        [
            ("Voice model (train_voice_model.py): ", True),
            (
                "Generates labeled feature rows, trains Random Forest and Gradient Boosting "
                "(MultiOutputRegressor), and saves the better average R² model as "
                "voice_fluency_model.pkl. If the pickle is missing, /analyze-voice falls back "
                "to a heuristic scorer.",
                False,
            ),
        ],
    )
    add_mixed(
        doc,
        [
            ("Camera model: ", True),
            (
                "OpenCV YuNet (face_detection_yunet_2023mar.onnx). Landmarks drive looking-away, "
                "looking-down, smile proxy, and fidgeting. Face count and brightness drive malpractice flags.",
                False,
            ),
        ],
    )

    add_table(
        doc,
        ["Module", "Input", "Algorithm / method", "Output"],
        [
            [
                "Answer evaluation (browser)",
                "Typed or spoken text + question spec",
                "Keyword aliases + regex patterns",
                "Score = 50% keywords + 35% patterns + 15% length",
            ],
            [
                "Camera analysis (Flask)",
                "JPEG webcam frame",
                "YuNet ONNX + landmark heuristics",
                "Confidence, nervousness, cheating risk, integrity",
            ],
            [
                "Voice analysis (Flask)",
                "Transcript, duration, confidence, keyword/pattern scores",
                "Random Forest or Gradient Boosting regressor",
                "fluency, accuracy, accent/clarity (0–100)",
            ],
        ],
        col_widths=[4.0, 4.2, 4.6, 4.0],
    )
    add_caption(doc, "Table 5. ML and scoring modules")

    add_heading_custom(doc, "2.7 Training datasets and model files", 2)
    add_body(
        doc,
        "CareerForge trains two sklearn models from CSV files in ml_service/. "
        "Face detection is not trained here: YuNet is a pretrained OpenCV ONNX file. "
        "There are no wav/mp3 corpora and no labeled webcam image sets. "
        "See also ml_service/TRAINING_DATA.md.",
    )
    add_table(
        doc,
        ["File", "Rows", "Kind", "Used by", "Output"],
        [
            [
                "skill_assessment_data.csv",
                "20",
                "Hand-written labels",
                "train_model.py",
                "skill_assessment_model.pkl, label_encoder.pkl",
            ],
            [
                "voice_fluency_data.csv",
                "480",
                "Synthetic features (seed 42)",
                "train_voice_model.py",
                "voice_fluency_model.pkl, voice_fluency_meta.pkl",
            ],
            [
                "face_detection_yunet_2023mar.onnx",
                "—",
                "Pretrained YuNet",
                "emotion_analyzer.py",
                "Not trained in this repo",
            ],
        ],
        col_widths=[4.2, 1.6, 3.4, 3.4, 4.2],
    )
    add_caption(doc, "Table 6. Training data and model artifacts")
    add_mixed(
        doc,
        [
            ("Skill data: ", True),
            (
                "20 rows. Features technical_score (45–95), communication_score (50–92), "
                "problem_solving (42–96). Labels Strong 7, Average 7, Weak 6. "
                "RandomForestClassifier with 100 trees, 80/20 split.",
                False,
            ),
        ],
    )
    add_mixed(
        doc,
        [
            ("Voice data: ", True),
            (
                "480 synthetic rows from five profiles (fluent 22%, average 28%, hesitant 20%, "
                "rushed 15%, inaccurate 15%). Eleven input features plus fluency_score and "
                "accuracy_score labels from teacher functions plus noise. Random Forest and "
                "Gradient Boosting compete; the higher average R² model is saved.",
                False,
            ),
        ],
    )
    add_body(
        doc,
        "Shipped pickles: skill_assessment_model.pkl, label_encoder.pkl, "
        "voice_fluency_model.pkl, voice_fluency_meta.pkl. Retrain with "
        "py -3 train_model.py and py -3 train_voice_model.py from ml_service/.",
    )

    # 3
    add_heading_custom(doc, "3. Sequence flows", 1)
    add_body(
        doc,
        "The following tables are sequence diagrams in document form. Read each row top to bottom "
        "as a message from one participant to another.",
    )

    add_heading_custom(doc, "3.1 End-to-end student journey", 2)
    add_sequence_table(
        doc,
        [
            ("Student", "React app", "Email/password or Google sign-in"),
            ("React app", "Express :5000", "POST /api/auth/login or /api/auth/google"),
            ("Express", "MySQL or JSON", "Upsert / load user"),
            ("Express", "React app", "JWT + user"),
            ("Student", "React app", "Open practice, courses, or SkillBridge"),
            ("React app", "Express", "GET /api/courses"),
            ("Express", "React app", "catalog.json (10 × 10)"),
            ("Student", "React app", "Select company and start mock"),
            ("React / Flask", "React / Flask", "Four-round interview loop"),
            ("React app", "Express", "POST /api/interview/results"),
            ("Express", "MySQL or JSON", "Save session"),
            ("React app", "React app", "Show /result: readiness, chance %, weak areas"),
        ],
    )
    add_caption(doc, "Figure 5. Sequence — student journey")
    add_body(doc, "The course list is served from src/data/catalog.json, not from a MySQL catalog table.")

    add_heading_custom(doc, "3.2 Answer scoring (keywords and patterns)", 2)
    add_sequence_table(
        doc,
        [
            ("Student", "Answer textarea", "Type or dictate answer"),
            ("Student", "MockInterview UI", "Submit answer"),
            ("MockInterview UI", "evaluateAnswer.js", "Send text + current question spec"),
            ("evaluateAnswer.js", "evaluateAnswer.js", "Match keyword aliases"),
            ("evaluateAnswer.js", "evaluateAnswer.js", "Test regex patterns"),
            ("evaluateAnswer.js", "MockInterview UI", "Keyword %, pattern %, length, combined score"),
            ("MockInterview UI", "MockInterview UI", "Store round; set next-round difficulty from last score"),
            ("MockInterview UI", "MockInterview UI", "Next of 4 rounds or finish"),
        ],
    )
    add_caption(doc, "Figure 6. Sequence — keyword and pattern matching")
    add_body(
        doc,
        "Keyword groups allow aliases (for example mutable also matches “can be changed”). "
        "Patterns check structure: definition phrasing, comparison words, examples, and prevention methods. "
        "Difficulty of the next round is foundation / applied / corporate from the last score — not a live "
        "Easy → Medium → Hard unlock while typing.",
    )

    add_heading_custom(doc, "3.3 Camera emotion and integrity", 2)
    add_sequence_table(
        doc,
        [
            ("Student", "Webcam", "Allow camera permission"),
            ("MockInterview", "Webcam", "getUserMedia (video)"),
            ("MockInterview", "Canvas", "Loop every ~1.4 seconds: capture JPEG"),
            ("MockInterview", "Vite proxy /ml", "POST /ml/analyze-emotion"),
            ("Vite proxy", "Flask emotion_analyzer", "POST /analyze-emotion"),
            ("Flask", "Flask", "YuNet detect faces and 5-point landmarks"),
            ("Flask", "Flask", "Score confidence, nervousness, cheating risk"),
            ("Flask", "MockInterview", "JSON + optional face bounding box"),
            ("MockInterview", "MockInterview", "Update meters and overlay"),
            ("MockInterview", "Result state", "Average samples into result.behavior"),
        ],
    )
    add_caption(doc, "Figure 7. Sequence — camera analysis")

    add_body(doc, "Integrity labels:")
    add_bullet(doc, " one frontal face, looking at the camera", "Clear —")
    add_bullet(doc, " looking away, fidgeting, moderate risk", "Warning —")
    add_bullet(doc, " no face, covered camera, or multiple people", "Flagged —")

    add_heading_custom(doc, "3.4 Voice fluency and accuracy", 2)
    add_sequence_table(
        doc,
        [
            ("Student", "MockInterview", "Start Voice Input"),
            ("MockInterview", "Web Speech API", "Start recognition"),
            ("Web Speech API", "MockInterview", "Transcript + confidence"),
            ("MockInterview", "Keyword/pattern scorer", "Score spoken text vs question"),
            ("MockInterview", "Flask /analyze-voice", "Transcript, duration, keyword/pattern scores"),
            ("Flask", "Flask", "extract_features()"),
            ("Flask", "Flask", "If model loaded: Random Forest / Gradient Boosting predict"),
            ("Flask", "Flask", "Else: heuristic_scores()"),
            ("Flask", "MockInterview", "Fluency, accuracy, accent/clarity, notes, algorithm name"),
            ("MockInterview", "MockInterview", "Show voice meters and attach scores to the answer"),
        ],
    )
    add_caption(doc, "Figure 8. Sequence — voice fluency, accuracy, and clarity")
    add_mixed(
        doc,
        [
            ("Fluency ", True),
            (
                "is driven by words per minute, filler ratio, unique-word ratio, repetition, pauses, "
                "and recognition confidence. ",
                False,
            ),
            ("Accuracy ", True),
            (
                "is driven by keyword coverage, pattern coverage, and speech-recognition confidence "
                "(how clearly the words were heard). ",
                False,
            ),
            ("Accent/clarity ", True),
            (
                "is an intelligibility estimate. It is not a linguistic accent identifier "
                "(it does not label “Indian English” vs “US English”).",
                False,
            ),
        ],
    )

    add_heading_custom(doc, "3.5 Persist interview session", 2)
    add_body(
        doc,
        "After the four rounds (or End interview), the result page posts a session. Express writes "
        "MySQL when the pool is up; otherwise it writes backend/data/sessions.json.",
    )
    add_sequence_table(
        doc,
        [
            ("React", "Express", "POST /api/interview/results (JWT)"),
            ("Express", "MySQL or JSON", "Save session for this user"),
            ("Store", "Express", "id"),
            ("Express", "React", "Saved confirmation"),
        ],
    )
    add_caption(doc, "Figure 9. Sequence — save interview session")

    # 4
    add_heading_custom(doc, "4. Usage — run on another computer", 1)
    add_body(
        doc,
        "Share the project folder (USB, zip, or git clone). The other PC must install Node.js and Python, "
        "then run the commands below. Three terminals stay open while the app is used. MySQL is optional: "
        "if it is missing, login and saved interviews still work using backend/data JSON files. "
        "Google Sign-In is optional.",
    )

    add_heading_custom(doc, "4.1 Software to install first (once per computer)", 2)
    add_numbered(doc, "Node.js LTS (20 or 22) from https://nodejs.org — this also installs npm. Restart the terminal after install.")
    add_numbered(doc, "Python 3.10+ from https://www.python.org/downloads/ — on Windows tick Add python.exe to PATH.")
    add_numbered(doc, "Google Chrome or Microsoft Edge (needed for Voice Input).")
    add_numbered(doc, "Optional: MySQL Server if you want a real database instead of JSON files.")
    add_body(doc, "Check that the tools work:")
    add_code(doc, "node -v\nnpm -v\npy -3 --version")
    add_body(
        doc,
        "On macOS / Linux use python3 --version instead of py -3 --version. If node or py is not found, "
        "close the terminal and open a new one, or reinstall with PATH enabled.",
    )

    add_heading_custom(doc, "4.2 What to copy (and what to leave out)", 2)
    add_body(
        doc,
        "Copy the whole ai_careerforge folder, including src/, backend/ (except secrets), ml_service/, "
        "and ml_service/models/face_detection_yunet_2023mar.onnx.",
    )
    add_table(
        doc,
        ["Leave out", "Why"],
        [
            ["node_modules/ and backend/node_modules/", "Too large; run npm install instead"],
            ["backend/.env", "Contains passwords; create a new file from .env.example"],
            ["backend/data/", "Local user accounts on the original machine"],
            ["dist/", "Production build output"],
            ["__pycache__/", "Python cache"],
        ],
        col_widths=[7.0, 9.8],
    )
    add_caption(doc, "Table 6. Files not to copy to the other computer")

    add_heading_custom(doc, "4.3 Windows — all commands (PowerShell)", 2)
    add_body(doc, "Open PowerShell. Replace the first cd with the real folder path on that PC.")
    add_mixed(doc, [("A. Install libraries (once)", True)])
    add_code(
        doc,
        "cd \"C:\\Users\\YOUR_NAME\\Desktop\\ai_careerforge\"\n"
        "\n"
        "node -v\n"
        "npm -v\n"
        "py -3 --version\n"
        "\n"
        "npm install\n"
        "\n"
        "cd backend\n"
        "npm install\n"
        "Copy-Item .env.example .env\n"
        "cd ..\n"
        "\n"
        "cd ml_service\n"
        "py -3 -m pip install --upgrade pip\n"
        "py -3 -m pip install -r requirements.txt\n"
        "py -3 train_model.py\n"
        "py -3 train_voice_model.py\n"
        "cd ..",
    )
    add_body(
        doc,
        "Edit backend\\.env in Notepad if MySQL is installed (set DB_PASSWORD). If MySQL is not installed, "
        "leave .env as copied — the API will print Accounts: file store and still work.",
    )
    add_mixed(doc, [("B. Start the app (every time) — three windows. Keep all three running.", True)])
    add_body(doc, "Terminal 1 — Flask ML (http://127.0.0.1:5001):")
    add_code(
        doc,
        "cd \"C:\\Users\\YOUR_NAME\\Desktop\\ai_careerforge\\ml_service\"\n"
        "py -3 app.py",
    )
    add_body(doc, "Terminal 2 — Express API (http://127.0.0.1:5000):")
    add_code(
        doc,
        "cd \"C:\\Users\\YOUR_NAME\\Desktop\\ai_careerforge\\backend\"\n"
        "node server.js",
    )
    add_body(doc, "Terminal 3 — React UI (http://127.0.0.1:5173):")
    add_code(
        doc,
        "cd \"C:\\Users\\YOUR_NAME\\Desktop\\ai_careerforge\"\n"
        "npm run dev",
    )
    add_mixed(doc, [("C. Open the app in Chrome", True)])
    add_code(doc, "http://127.0.0.1:5173")
    add_body(
        doc,
        "Use 127.0.0.1, not localhost, if you later add Google Sign-In (Google treats them as different sites).",
    )

    add_heading_custom(doc, "4.4 macOS / Linux — all commands", 2)
    add_mixed(doc, [("A. Install libraries (once)", True)])
    add_code(
        doc,
        "cd ~/Desktop/ai_careerforge\n"
        "\n"
        "node -v\n"
        "npm -v\n"
        "python3 --version\n"
        "\n"
        "npm install\n"
        "\n"
        "cd backend\n"
        "npm install\n"
        "cp .env.example .env\n"
        "cd ..\n"
        "\n"
        "cd ml_service\n"
        "python3 -m pip install --upgrade pip\n"
        "python3 -m pip install -r requirements.txt\n"
        "python3 train_model.py\n"
        "python3 train_voice_model.py\n"
        "cd ..",
    )
    add_mixed(doc, [("B. Start the app (every time) — three terminals", True)])
    add_body(doc, "Terminal 1:")
    add_code(doc, "cd ~/Desktop/ai_careerforge/ml_service\npython3 app.py")
    add_body(doc, "Terminal 2:")
    add_code(doc, "cd ~/Desktop/ai_careerforge/backend\nnode server.js")
    add_body(doc, "Terminal 3:")
    add_code(doc, "cd ~/Desktop/ai_careerforge\nnpm run dev")
    add_body(doc, "Open http://127.0.0.1:5173")

    add_heading_custom(doc, "4.5 Optional MySQL", 2)
    add_body(doc, "Only if you want accounts in MySQL instead of JSON files.")
    add_code(doc, "CREATE DATABASE ai_interview_career_support_db;")
    add_body(doc, "Then set these in backend/.env (never share this file):")
    add_code(
        doc,
        "DB_HOST=localhost\n"
        "DB_USER=root\n"
        "DB_PASSWORD=YOUR_MYSQL_PASSWORD\n"
        "DB_NAME=ai_interview_career_support_db\n"
        "JWT_SECRET=any-long-random-string\n"
        "PORT=5000\n"
        "ML_ORIGIN=http://127.0.0.1:5001\n"
        "GOOGLE_CLIENT_ID=",
    )
    add_body(doc, "Tables are created automatically on API start. Restart node server.js after editing .env.")

    add_heading_custom(doc, "4.6 Optional Google Sign-In", 2)
    add_numbered(doc, "Google Cloud Console → APIs & Services → Credentials → Create OAuth client ID → Web application.")
    add_numbered(doc, "Authorized JavaScript origins: http://127.0.0.1:5173 and http://localhost:5173 and http://127.0.0.1:5000.")
    add_numbered(doc, "Paste the client ID into GOOGLE_CLIENT_ID= in backend/.env.")
    add_numbered(doc, "Restart node server.js.")
    add_body(doc, "Leave GOOGLE_CLIENT_ID empty to use email/password only.")

    add_heading_custom(doc, "4.7 Check it is running", 2)
    add_body(doc, "With the three terminals open, run in a fourth window:")
    add_code(doc, "curl http://127.0.0.1:5001/health\ncurl http://127.0.0.1:5000/api/health")
    add_body(doc, "You should see JSON with ok: true. Then open http://127.0.0.1:5173 and register a new account.")
    add_table(
        doc,
        ["Problem", "Fix"],
        [
            ["node is not recognized", "Reinstall Node.js LTS; close and reopen the terminal"],
            ["py / python is not recognized", "Reinstall Python with Add to PATH; use py -3 on Windows"],
            ["Port 5173 / 5000 / 5001 already in use", "Close the old CareerForge windows and start again"],
            ["Camera or voice offline", "Terminal 1 (py -3 app.py) must be running; allow camera/mic in Chrome"],
            ["Accounts: file store in the API window", "Normal when MySQL is off — login still works"],
            ["Page loads but login fails", "Confirm Terminal 2 is running; use http://127.0.0.1:5173"],
        ],
        col_widths=[6.4, 10.4],
    )
    add_caption(doc, "Table 7. Common setup problems")

    add_heading_custom(doc, "4.8 How to use the product", 2)
    add_numbered(doc, "Register or sign in with email/password, or Google if GOOGLE_CLIENT_ID is set.")
    add_numbered(doc, "Dashboard: Practice, course path, Explore, Guide, or SkillBridge (/readiness).")
    add_numbered(doc, "Practice: pick course → company → open round resources → start mock; or Courses → company → preparation → mock.")
    add_numbered(doc, "Complete four adaptive rounds (aptitude, coding, technical, HR — names follow the course track).")
    add_numbered(doc, "Answer box: type a full answer. Scoring runs on submit (keywords 50% + patterns 35% + length 15%).")
    add_numbered(doc, "Camera: allow access. Watch confidence, nervousness, and malpractice risk. Sit in a bright room, one person in frame, facing the camera.")
    add_numbered(doc, "Voice: use Voice Input in Chrome. Flask scores fluency, accuracy, and accent/clarity when the ML service is up.")
    add_numbered(doc, "Next-round difficulty follows the last round’s score (foundation / applied / corporate).")
    add_numbered(doc, "Result page: company readiness, chance of advancing this screen, accent/clarity, weak areas. Session is saved to the signed-in account.")

    add_heading_custom(doc, "4.9 Main API map", 2)
    add_table(
        doc,
        ["Method", "Path", "Purpose"],
        [
            ["GET", "/api/health", "API health"],
            ["GET", "/api/auth/config", "Public Google client id (may be empty)"],
            ["POST", "/api/auth/register", "Email/password register"],
            ["POST", "/api/auth/login", "Email/password login"],
            ["POST", "/api/auth/google", "Google ID-token login"],
            ["GET", "/api/auth/me", "Current user (JWT)"],
            ["GET", "/api/courses", "10 courses from catalog.json"],
            ["GET", "/api/courses/:id/companies", "10 companies for a course"],
            ["GET", "/api/companies/:id/requirements", "Company requirements"],
            ["GET", "/api/requirements/detect", "Hiring-cycle detector (static snapshots)"],
            ["POST", "/api/interview/results", "Save interview session (JWT)"],
            ["GET", "/api/me/results", "Saved sessions for the signed-in user"],
            ["GET", "http://127.0.0.1:5001/", "ML health"],
            ["POST", "http://127.0.0.1:5001/predict", "Skill level from three scores"],
            ["POST", "http://127.0.0.1:5001/analyze-emotion", "Camera frame"],
            ["POST", "http://127.0.0.1:5001/analyze-voice", "Fluency, accuracy, accent/clarity"],
        ],
        col_widths=[2.2, 8.4, 6.2],
    )
    add_caption(doc, "Table 8. HTTP endpoints")

    # 5
    add_heading_custom(doc, "5. Pros and cons", 1)

    add_heading_custom(doc, "5.1 Pros", 2)
    add_bullet(doc, " from Practice / SkillBridge to a four-round scored result, not a single quiz page.", "End-to-end loop —")
    add_bullet(doc, " keyword and pattern rules are documented; scores appear after submit.", "Explainable answer scoring —")
    add_bullet(doc, " next-round difficulty follows the last score (foundation / applied / corporate).", "Adaptive rounds —")
    add_bullet(doc, " OpenCV and sklearn stay in Python; the UI stays in React.", "Separate ML service —")
    add_bullet(doc, " extra faces, empty frame, covered camera, looking down.", "Integrity signals —")
    add_bullet(doc, " Random Forest vs Gradient Boosting, heuristic fallback if the pickle is absent.", "Voice model is trainable —")
    add_bullet(doc, " JWT, optional Google Sign-In, MySQL or JSON files.", "Auth with fallback storage —")
    add_bullet(doc, " one origin for the browser via Vite /api and /ml.", "Dev-friendly proxy —")

    add_heading_custom(doc, "5.2 Cons", 2)
    add_bullet(doc, " Keyword lists can miss a valid unusual answer; camera heuristics can flag looking at notes when the student is only thinking.", "Not a substitute for a human interviewer.")
    add_bullet(doc, " Confidence and nervousness are behavioral estimates, not clinical emotion recognition.", "YuNet cues are approximate.")
    add_bullet(doc, " Accents, noise, and Chrome-only Speech API limit quality.", "Voice accuracy depends on the browser transcript.")
    add_bullet(doc, " It is intelligibility, not a linguistic accent model.", "Accent/clarity is not an accent classifier.")
    add_bullet(doc, " Generated feature rows mean the regressor learns a designed scoring policy, not a large corpus of real interviews.", "Voice training data is synthetic.")
    add_bullet(doc, " Diffs static snapshots in companySignals.js; it does not scrape live job boards.", "Hiring-cycle detector is not live.")
    add_bullet(doc, " If Flask is down, camera/voice analysis show offline errors while typed scoring still works.", "Three processes must run together.")
    add_bullet(doc, " Needs a real OAuth client ID in .env; empty id hides the Google button.", "Google Sign-In is optional.")

    add_table(
        doc,
        ["Area", "Advantage", "Limitation"],
        [
            ["Answer scoring", "Visible keywords and patterns", "Can miss valid wording not in the list"],
            ["Camera", "Live integrity and presence cues", "Not medical emotion detection"],
            ["Voice ML", "Trained RF/GBR with fallback", "Synthetic labels; depends on browser STT"],
            ["Architecture", "Clear split of UI / API / ML + block diagrams", "Three services to start"],
        ],
        col_widths=[3.6, 6.6, 6.6],
    )
    add_caption(doc, "Table 9. Pros and cons at a glance")

    # 6
    add_heading_custom(doc, "6. Project structure", 1)
    add_code(
        doc,
        "ai_careerforge/\n"
        "  src/\n"
        "    App.jsx                 # React Router\n"
        "    data/\n"
        "      catalog.json          # 10 courses × 10 companies\n"
        "      interviewBank.js      # mock questions\n"
        "      practiceResources.js  # round resource URLs\n"
        "      skillMaps.js          # SkillBridge chains\n"
        "      companySignals.js     # hiring-cycle snapshots\n"
        "    pages/                  # Login, Dashboard, Practice, interview, Result, ReadinessHub\n"
        "    utils/\n"
        "      evaluateAnswer.js     # keywords + regex patterns\n"
        "      adaptiveInterview.js  # 4 rounds, difficulty, conversionChance\n"
        "      explainSkills.js      # result pack\n"
        "      accentScore.js        # client fallback if Flask is down\n"
        "  backend/\n"
        "    server.js               # Express REST + /ml proxy in production\n"
        "    auth.js / accounts.js   # JWT, Google upsert\n"
        "    db.js                   # MySQL pool from .env\n"
        "    data/                   # JSON fallback (gitignored)\n"
        "  ml_service/\n"
        "    app.py                  # Flask: predict, analyze-emotion, analyze-voice\n"
        "    emotion_analyzer.py     # YuNet + integrity heuristics\n"
        "    voice_features.py       # fluency / accuracy / accent features\n"
        "    train_model.py          # skill Random Forest\n"
        "    train_voice_model.py    # voice Random Forest / GBR\n"
        "    TRAINING_DATA.md        # dataset and model inventory\n"
        "    skill_assessment_data.csv\n"
        "    voice_fluency_data.csv\n"
        "    skill_assessment_model.pkl / label_encoder.pkl\n"
        "    voice_fluency_model.pkl / voice_fluency_meta.pkl\n"
        "    models/                 # YuNet ONNX (pretrained)\n"
        "  docs/figures/            # architecture and block-diagram PNGs\n"
        "  DOCUMENTATION.md\n"
        "  CareerForge_Documentation.docx",
    )

    # 7
    add_heading_custom(doc, "7. Limitations and future work", 1)
    add_bullet(doc, "Replace synthetic voice labels with recorded mock interviews and human fluency/accuracy ratings.")
    add_bullet(doc, "Use a dedicated speech model (Whisper / wav2vec) instead of browser transcription.")
    add_bullet(doc, "Train a real accent classifier only with labeled speech; do not treat the current score as ethnicity or region detection.")
    add_bullet(doc, "Prefer MySQL in production and keep JSON only as a local fallback.")
    add_bullet(doc, "Expand live job-board ingestion if the requirement detector should reflect the market, not snapshots.")
    add_bullet(doc, "Put a reverse proxy (nginx) in front of the built SPA, Express, and Flask for a single production origin.")

    # 8
    add_heading_custom(doc, "8. Summary", 1)
    add_body(
        doc,
        "CareerForge is a three-process mock-interview system: a React SPA (Vite :5173), an Express API "
        "(:5000) with JWT/Google auth and a 10×10 catalog, and a Flask ML service (:5001) for camera and voice. "
        "Product blocks are prepare (practice, SkillBridge, resources), mock (four adaptive rounds), scoring "
        "(text, camera, voice), and output (readiness, chance of advancing, weak areas). The models are "
        "practice tools, not production proctoring or medical diagnostics.",
    )

    out = Path(r"c:\Users\bhkp\Desktop\ai interview\ai_careerforge\CareerForge_Documentation.docx")
    try:
        doc.save(str(out))
    except PermissionError:
        out = Path(r"c:\Users\bhkp\Desktop\ai interview\ai_careerforge\CareerForge_Documentation_updated.docx")
        doc.save(str(out))
        print("CareerForge_Documentation.docx is open in Word; wrote:")
    print(out)


if __name__ == "__main__":
    build()
